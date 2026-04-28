import fitz  # PyMuPDF
import docx
import io
import re
from typing import Dict, List, Set

# ── URL patterns for plain-text extraction ──────────────────────────────
GITHUB_PROFILE_RE = re.compile(
    r'https?://(?:www\.)?github\.com/([A-Za-z0-9_.-]+)(?:/)?(?:\s|$|[),\]\}>])',
    re.IGNORECASE
)
GITHUB_REPO_RE = re.compile(
    r'https?://(?:www\.)?github\.com/([A-Za-z0-9_.-]+)/([A-Za-z0-9_.-]+)',
    re.IGNORECASE
)
LEETCODE_RE = re.compile(
    r'https?://(?:www\.)?leetcode\.com/(?:u/)?([A-Za-z0-9_-]+)',
    re.IGNORECASE
)
LINKEDIN_RE = re.compile(
    r'https?://(?:www\.)?linkedin\.com/in/([A-Za-z0-9_-]+)',
    re.IGNORECASE
)
# Generic URL pattern to capture any http(s) link
URL_RE = re.compile(r'https?://[^\s<>"\')\]},]+', re.IGNORECASE)


class Parser:
    def parse(self, file_bytes: bytes, ext: str):
        if ext == 'pdf':
            text = self._parse_pdf(file_bytes)
        elif ext == 'docx':
            text = self._parse_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported extension: {ext}")
            
        sections = self._extract_sections(text)
        sentences = self._split_into_sentences(text)
        links = self._extract_links(file_bytes, ext, text)
        
        return {
            "rawText": text,
            "sections": sections,
            "sentences": sentences,
            "links": links
        }

    # ── Link extraction (annotations + regex) ─────────────────────────────

    def _extract_links(self, file_bytes: bytes, ext: str, raw_text: str) -> Dict[str, List[str]]:
        """
        Extract GitHub, LeetCode, LinkedIn URLs from:
          1) Clickable hyperlink annotations embedded in the PDF/DOCX
          2) Plain-text URLs found via regex in the raw extracted text
        """
        raw_urls: Set[str] = set()

        # ── Strategy 1: annotation / hyperlink objects ──
        try:
            if ext == 'pdf':
                raw_urls.update(self._pdf_annotation_links(file_bytes))
            elif ext == 'docx':
                raw_urls.update(self._docx_hyperlinks(file_bytes))
        except Exception as e:
            print(f"[Parser] Warning: annotation link extraction failed: {e}")

        # ── Strategy 2: regex on raw text ──
        raw_urls.update(URL_RE.findall(raw_text))

        # ── Classify ──
        github_profiles: List[str] = []
        repos: List[str] = []
        leetcode: List[str] = []
        linkedin: List[str] = []

        seen_github_users: Set[str] = set()
        seen_repos: Set[str] = set()

        for url in raw_urls:
            url_clean = url.rstrip('/').rstrip('.,;:')

            # GitHub repo URLs (user/repo) — check first (more specific)
            repo_match = GITHUB_REPO_RE.search(url_clean)
            if repo_match:
                user = repo_match.group(1).lower()
                repo = repo_match.group(2).lower()
                # Skip github.com meta pages
                if user not in ('topics', 'explore', 'settings', 'notifications', 'login', 'signup', 'features'):
                    repo_key = f"{user}/{repo}"
                    if repo_key not in seen_repos:
                        seen_repos.add(repo_key)
                        repos.append(url_clean)
                    # Also add profile URL if not seen
                    if user not in seen_github_users:
                        seen_github_users.add(user)
                        github_profiles.append(f"https://github.com/{repo_match.group(1)}")
                continue

            # GitHub profile URLs (just user, no repo)
            profile_match = GITHUB_PROFILE_RE.search(url_clean + ' ')  # trailing space for regex
            if profile_match:
                user = profile_match.group(1).lower()
                if user not in ('topics', 'explore', 'settings', 'notifications', 'login', 'signup', 'features'):
                    if user not in seen_github_users:
                        seen_github_users.add(user)
                        github_profiles.append(url_clean)
                continue

            # LeetCode
            lc_match = LEETCODE_RE.search(url_clean)
            if lc_match:
                username = lc_match.group(1).lower()
                if username not in ('problemset', 'contest', 'discuss', 'explore', 'playground'):
                    leetcode.append(url_clean)
                continue

            # LinkedIn
            li_match = LINKEDIN_RE.search(url_clean)
            if li_match:
                linkedin.append(url_clean)

        # Deduplicate while preserving order
        return {
            "github":   list(dict.fromkeys(github_profiles)),
            "leetcode": list(dict.fromkeys(leetcode)),
            "linkedin": list(dict.fromkeys(linkedin)),
            "repos":    list(dict.fromkeys(repos))
        }

    def _pdf_annotation_links(self, file_bytes: bytes) -> List[str]:
        """Extract clickable URI links from PDF page annotations."""
        urls = []
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            for link in page.get_links():
                if link.get("kind") == 2:  # kind 2 = external URI
                    uri = link.get("uri", "")
                    if uri and uri.startswith("http"):
                        urls.append(uri)
        return urls

    def _docx_hyperlinks(self, file_bytes: bytes) -> List[str]:
        """Extract hyperlinks from DOCX paragraphs."""
        urls = []
        doc = docx.Document(io.BytesIO(file_bytes))
        # Method 1: paragraph hyperlinks (python-docx >= 1.x)
        for para in doc.paragraphs:
            if hasattr(para, 'hyperlinks'):
                for hl in para.hyperlinks:
                    if hasattr(hl, 'url') and hl.url:
                        urls.append(hl.url)
        # Method 2: relationship-level fallback (older python-docx)
        if not urls:
            try:
                for rel in doc.part.rels.values():
                    if "hyperlink" in rel.reltype:
                        target = getattr(rel, 'target_ref', None)
                        if target and target.startswith("http"):
                            urls.append(target)
            except Exception:
                pass
        return urls

    # ── Existing methods ──────────────────────────────────────────────────

    def _parse_pdf(self, file_bytes: bytes) -> str:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return text

    def _parse_docx(self, file_bytes: bytes) -> str:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = "\n".join([p.text for p in doc.paragraphs])
        return text
        
    def _extract_sections(self, text: str) -> dict:
        # A simple regex-based heuristic section extractor
        sections = {
            "summary": "",
            "skills": "",
            "experience": "",
            "education": "",
            "projects": "",
            "certifications": "",
            "other": ""
        }
        
        current_section = "other"
        lines = text.split('\n')
        
        for line in lines:
            line_clean = line.strip().lower()
            if not line_clean:
                continue
                
            if re.match(r'^(summary|profile|about me)', line_clean):
                current_section = "summary"
            elif re.match(r'^(skills|technologies|core competencies)', line_clean):
                current_section = "skills"
            elif re.match(r'^(experience|employment|work history|professional experience)', line_clean):
                current_section = "experience"
            elif re.match(r'^(education|academic background)', line_clean):
                current_section = "education"
            elif re.match(r'^(projects|personal projects)', line_clean):
                current_section = "projects"
            elif re.match(r'^(certifications|licenses)', line_clean):
                current_section = "certifications"
            else:
                sections[current_section] += line + "\n"
                
        return sections

    def _split_into_sentences(self, text: str) -> list:
        # Simple heuristic split on punctuation
        sentences = re.split(r'(?<=[.!?]) +', text.replace('\n', ' '))
        return [s.strip() for s in sentences if len(s.strip()) > 5]

parser = Parser()
